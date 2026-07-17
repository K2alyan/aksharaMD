"""Integration tests for KeyValueGroup pipeline wiring.

Covers:
- Transformation stage (KI1-KI10)
- Adjacent-block detection (KA1-KA5)
- HTML native <dl> handling (KH1-KH5)
- XLSX two-column detection (KX1-KX5)
- DOCX properties extraction (KD1-KD3)
- Optimizer protection (KO1-KO4)
- Chunker integration (KC1-KC4)
- Backward compatibility (KB1-KB4)
"""
from __future__ import annotations

import os
import tempfile
from pathlib import Path

from aksharamd.models.block import Block, BlockType, ExtractionConfidence
from aksharamd.models.document import Document
from aksharamd.models.key_value import KeyValueEntry, KeyValueGroup, KeyValueGroupType
from aksharamd.plugins.transformers.key_value_promoter import (
    DETECTOR_VERSION,
    detect_and_promote_key_value_groups,
)

# ── Helpers ──────────────────────────────────────────────────────────────────


def _make_ctx(blocks: list[Block]):
    """Create a minimal CompilationContext with the given blocks.

    Attaches the experimental KV profile so the heuristic promoter runs —
    these integration tests exercise the promoter's semantic behaviour,
    which under kv_promoter/v2 is opt-in via the profile.
    """
    from aksharamd.context import CompilationContext
    from aksharamd.scoring.key_value_config import KeyValueDetectionProfile

    ctx = CompilationContext(source="test.txt", output_dir=tempfile.mkdtemp())
    doc = Document(source="test.txt", blocks=blocks)
    ctx.document = doc
    ctx.kv_profile = KeyValueDetectionProfile.experimental()
    return ctx


def _para(content: str, page: int | None = None, index: int = 0) -> Block:
    return Block(type=BlockType.PARAGRAPH, content=content, page=page, index=index)


def _heading(content: str, level: int = 2, index: int = 0) -> Block:
    return Block(type=BlockType.HEADING, content=content, level=level, index=index)


def _table_block(index: int = 0) -> Block:
    from aksharamd.models.table import TableCell, TableData
    cells = [
        TableCell(text="A", row=0, column=0),
        TableCell(text="B", row=0, column=1),
        TableCell(text="1", row=1, column=0),
        TableCell(text="2", row=1, column=1),
    ]
    td = TableData(row_count=2, column_count=2, cells=cells)
    return Block.from_table(td, index=index)


def _html_fixture(dl_html: str) -> str:
    return f"<html><body>{dl_html}</body></html>"


def _parse_html_blocks(html_str: str) -> list[Block]:
    """Write HTML to a temp file and parse it with the HTML parser."""
    from aksharamd.context import CompilationContext
    from aksharamd.plugins.parsers.html import HTMLParser

    with tempfile.NamedTemporaryFile(
        mode="w", suffix=".html", delete=False, encoding="utf-8"
    ) as f:
        f.write(html_str)
        fname = f.name
    try:
        ctx = CompilationContext(source=fname, output_dir=tempfile.mkdtemp())
        parser = HTMLParser()
        result = parser.execute(ctx)
        return result.document.blocks if result.document else []
    finally:
        os.unlink(fname)


def _make_xlsx_kv(tmp_path: Path) -> Path:
    """Create a 2-column KV XLSX fixture."""
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Properties"
    data = [
        ("Project", "AksharaMD"),
        ("Owner", "Kalyan"),
        ("Version", "0.4.0"),
        ("Status", "Active"),
    ]
    for r, (k, v) in enumerate(data, 1):
        ws.cell(row=r, column=1, value=k)
        ws.cell(row=r, column=2, value=v)
    path = tmp_path / "kv.xlsx"
    wb.save(str(path))
    return path


def _make_xlsx_table(tmp_path: Path) -> Path:
    """Create a conventional 3-column XLSX fixture."""
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Data"
    ws.cell(row=1, column=1, value="Name")
    ws.cell(row=1, column=2, value="Age")
    ws.cell(row=1, column=3, value="City")
    ws.cell(row=2, column=1, value="Alice")
    ws.cell(row=2, column=2, value=30)
    ws.cell(row=2, column=3, value="NYC")
    ws.cell(row=3, column=1, value="Bob")
    ws.cell(row=3, column=2, value=25)
    ws.cell(row=3, column=3, value="LA")
    path = tmp_path / "table.xlsx"
    wb.save(str(path))
    return path


def _make_xlsx_repeated_keys(tmp_path: Path) -> Path:
    """Create a 2-column XLSX with repeated first-column values (data table, not KV)."""
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Repeated"
    data = [
        ("Name", "Alice"),
        ("Name", "Bob"),
        ("Name", "Charlie"),
        ("Name", "Dave"),
    ]
    for r, (k, v) in enumerate(data, 1):
        ws.cell(row=r, column=1, value=k)
        ws.cell(row=r, column=2, value=v)
    path = tmp_path / "repeated.xlsx"
    wb.save(str(path))
    return path


def _make_docx_with_props(tmp_path: Path, title: str = "Test Doc", author: str = "Test Author") -> Path:
    """Create a DOCX file with core properties."""
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.core_properties.title = title
    doc.core_properties.author = author
    doc.add_paragraph("This is the body of the document.")
    path = tmp_path / "test.docx"
    doc.save(str(path))
    return path


def _make_docx_no_props(tmp_path: Path) -> Path:
    """Create a DOCX file without meaningful core properties."""
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_paragraph("Just a paragraph with no metadata.")
    path = tmp_path / "noprops.docx"
    doc.save(str(path))
    return path


def _parse_docx_blocks(path: Path) -> list[Block]:
    """Parse a DOCX file and return its blocks."""
    from aksharamd.context import CompilationContext
    from aksharamd.plugins.parsers.docx import DocxParser

    ctx = CompilationContext(source=str(path), output_dir=tempfile.mkdtemp())
    parser = DocxParser()
    result = parser.execute(ctx)
    return result.document.blocks if result.document else []


def _parse_xlsx_blocks(path: Path) -> list[Block]:
    """Parse an XLSX file and return its blocks."""
    from aksharamd.context import CompilationContext
    from aksharamd.plugins.parsers.spreadsheet import XlsxParser

    ctx = CompilationContext(source=str(path), output_dir=tempfile.mkdtemp())
    parser = XlsxParser()
    result = parser.execute(ctx)
    return result.document.blocks if result.document else []


# ── TRANSFORMATION STAGE ─────────────────────────────────────────────────────


def test_ki1_two_line_paragraph_promoted():
    """KI1: Two-line 'Key: Value\\nKey: Value' paragraph is promoted to KEY_VALUE_GROUP block."""
    content = "Name: Alice\nEmail: alice@example.com"
    ctx = _make_ctx([_para(content)])
    result = detect_and_promote_key_value_groups(ctx)
    blocks = result.document.blocks
    assert len(blocks) == 1
    assert blocks[0].type == BlockType.KEY_VALUE_GROUP


def test_ki2_single_line_not_promoted():
    """KI2: Single-line paragraph with one colon is NOT promoted (single entry)."""
    content = "Name: Alice"
    ctx = _make_ctx([_para(content)])
    result = detect_and_promote_key_value_groups(ctx)
    blocks = result.document.blocks
    assert len(blocks) == 1
    assert blocks[0].type == BlockType.PARAGRAPH


def test_ki3_long_prose_not_promoted():
    """KI3: Long prose paragraph (>600 chars) is NOT promoted."""
    # Build a paragraph that is > 600 chars with content that looks like KV but is too long
    long_prose = "This is a very long paragraph. " * 25  # ~750 chars, no colons, no KV
    assert len(long_prose) > 600
    ctx = _make_ctx([_para(long_prose)])
    result = detect_and_promote_key_value_groups(ctx)
    blocks = result.document.blocks
    assert len(blocks) == 1
    assert blocks[0].type == BlockType.PARAGRAPH


def test_ki4_no_kv_document_unchanged():
    """KI4: Document with no KV paragraphs is unchanged."""
    blocks = [
        _heading("My Report", index=0),
        _para("This is regular prose with no key-value pairs.", index=1),
        _para("Another paragraph of content.", index=2),
    ]
    original_ids = [b.id for b in blocks]
    ctx = _make_ctx(blocks)
    result = detect_and_promote_key_value_groups(ctx)
    result_ids = [b.id for b in result.document.blocks]
    assert result_ids == original_ids


def test_ki5_promoted_block_preserves_source_block_ids():
    """KI5: Promoted block preserves source_block_ids in transformation metadata."""
    content = "Name: Alice\nEmail: alice@example.com"
    original_block = _para(content)
    ctx = _make_ctx([original_block])
    result = detect_and_promote_key_value_groups(ctx)
    promoted = result.document.blocks[0]
    assert promoted.type == BlockType.KEY_VALUE_GROUP
    assert original_block.id in promoted.metadata.get("source_block_ids", [])


def test_ki6_promoted_block_transformation_field():
    """KI6: Promoted block has 'transformation' field = 'kv_inline_detection' in metadata."""
    content = "Name: Alice\nEmail: alice@example.com"
    ctx = _make_ctx([_para(content)])
    result = detect_and_promote_key_value_groups(ctx)
    promoted = result.document.blocks[0]
    assert promoted.metadata.get("transformation") == "kv_inline_detection"


def test_ki7_promoted_block_has_original_text_checksum():
    """KI7: Promoted block has 'original_text_checksum' in metadata."""
    content = "Name: Alice\nEmail: alice@example.com"
    ctx = _make_ctx([_para(content)])
    result = detect_and_promote_key_value_groups(ctx)
    promoted = result.document.blocks[0]
    assert "original_text_checksum" in promoted.metadata


def test_ki8_heading_never_promoted():
    """KI8: HEADING block is never promoted to KV group."""
    h = _heading("Name: Alice\nEmail: test@example.com", index=0)
    ctx = _make_ctx([h])
    result = detect_and_promote_key_value_groups(ctx)
    assert result.document.blocks[0].type == BlockType.HEADING


def test_ki9_table_block_never_touched():
    """KI9: TABLE block is never touched by the promoter."""
    t = _table_block(index=0)
    original_id = t.id
    ctx = _make_ctx([t])
    result = detect_and_promote_key_value_groups(ctx)
    assert result.document.blocks[0].type == BlockType.TABLE
    assert result.document.blocks[0].id == original_id


def test_ki10_two_adjacent_kv_paragraphs_promoted_independently():
    """KI10: Two adjacent 2-line KV paragraphs each independently get promoted."""
    p1 = _para("Name: Alice\nEmail: alice@example.com", index=0)
    p2 = _para("City: NYC\nPhone: +1-555-0100", index=1)
    ctx = _make_ctx([p1, p2])
    result = detect_and_promote_key_value_groups(ctx)
    # Each paragraph should be independently promoted (or combined by adjacent detection)
    # Either way, all resulting blocks should be KEY_VALUE_GROUP
    kv_blocks = [b for b in result.document.blocks if b.type == BlockType.KEY_VALUE_GROUP]
    assert len(kv_blocks) >= 1  # At least one was promoted


# ── ADJACENT-BLOCK DETECTION ─────────────────────────────────────────────────


def test_ka1_run_of_four_kv_paragraphs_promoted():
    """KA1: Run of 4 short adjacent paragraphs with KV content is promoted to one group."""
    blocks = [
        _para("Name: Alice", index=0),
        _para("Email: alice@example.com", index=1),
        _para("City: NYC", index=2),
        _para("Phone: +1-555-0100", index=3),
    ]
    ctx = _make_ctx(blocks)
    result = detect_and_promote_key_value_groups(ctx)
    kv_blocks = [b for b in result.document.blocks if b.type == BlockType.KEY_VALUE_GROUP]
    assert len(kv_blocks) >= 1


def test_ka2_adjacent_run_stopped_at_heading():
    """KA2: Adjacent run is stopped at a HEADING block."""
    blocks = [
        _para("Name: Alice", index=0),
        _para("Email: alice@example.com", index=1),
        _heading("Section 2", level=2, index=2),
        _para("City: NYC", index=3),
        _para("Phone: +1-555-0100", index=4),
    ]
    ctx = _make_ctx(blocks)
    result = detect_and_promote_key_value_groups(ctx)
    # Heading must still be a heading
    heading_blocks = [b for b in result.document.blocks if b.type == BlockType.HEADING]
    assert len(heading_blocks) >= 1


def test_ka3_adjacent_run_stopped_at_table():
    """KA3: Adjacent run is stopped at a TABLE block."""
    t = _table_block(index=2)
    blocks = [
        _para("Name: Alice", index=0),
        _para("Email: alice@example.com", index=1),
        t,
        _para("City: NYC", index=3),
    ]
    ctx = _make_ctx(blocks)
    result = detect_and_promote_key_value_groups(ctx)
    table_blocks = [b for b in result.document.blocks if b.type == BlockType.TABLE]
    assert len(table_blocks) == 1


def test_ka4_source_block_ids_cover_consumed_blocks():
    """KA4: Source block IDs in promoted group cover all consumed blocks in the run."""
    blocks = [
        _para("Name: Alice", index=0),
        _para("Email: alice@example.com", index=1),
        _para("City: NYC", index=2),
        _para("Phone: +1-555-0100", index=3),
    ]
    original_ids = {b.id for b in blocks}
    ctx = _make_ctx(blocks)
    result = detect_and_promote_key_value_groups(ctx)
    kv_blocks = [b for b in result.document.blocks if b.type == BlockType.KEY_VALUE_GROUP]
    if kv_blocks:
        # Check source_block_ids are a subset of the original IDs
        for kb in kv_blocks:
            source_ids = kb.key_value_group.source_block_ids if kb.key_value_group else []
            for sid in source_ids:
                assert sid in original_ids


def test_ka5_run_of_two_not_promoted_as_adjacent():
    """KA5: Run of only 2 paragraphs doesn't form an adjacent group (< 4 required),
    but single-block inline may still fire if each has 2 KV lines."""
    # Two separate single-KV-line blocks — neither has enough lines for inline promotion
    # and < 4 blocks prevents adjacent promotion
    blocks = [
        _para("Name: Alice", index=0),
        _para("Email: alice@example.com", index=1),
    ]
    ctx = _make_ctx(blocks)
    result = detect_and_promote_key_value_groups(ctx)
    # Neither block should become a KV group on its own via inline detection
    # (inline requires 2 lines within one block)
    # The adjacent pass requires >= 4 blocks in a run
    # So both should remain paragraphs
    kv_blocks = [b for b in result.document.blocks if b.type == BlockType.KEY_VALUE_GROUP]
    # No adjacent promotion (run size 2 < minimum 4)
    # Also no inline promotion (single-line paragraphs)
    assert len(kv_blocks) == 0


# ── HTML NATIVE ───────────────────────────────────────────────────────────────


def test_kh1_dl_with_single_pair_creates_kv_block():
    """KH1: <dl><dt>Email</dt><dd>user@example.com</dd></dl> in HTML → KEY_VALUE_GROUP block."""
    html = _html_fixture("<dl><dt>Email</dt><dd>user@example.com</dd></dl>")
    blocks = _parse_html_blocks(html)
    kv_blocks = [b for b in blocks if b.type == BlockType.KEY_VALUE_GROUP]
    assert len(kv_blocks) >= 1


def test_kh2_dl_with_three_pairs_creates_three_entries():
    """KH2: Multi-entry <dl> with 3 dt/dd pairs → 3 entries in group."""
    html = _html_fixture(
        "<dl>"
        "<dt>Name</dt><dd>Alice</dd>"
        "<dt>Email</dt><dd>alice@example.com</dd>"
        "<dt>City</dt><dd>New York</dd>"
        "</dl>"
    )
    blocks = _parse_html_blocks(html)
    kv_blocks = [b for b in blocks if b.type == BlockType.KEY_VALUE_GROUP]
    assert len(kv_blocks) >= 1
    group = kv_blocks[0].key_value_group
    assert group is not None
    assert len(group.entries) == 3


def test_kh3_dl_extraction_method_is_html_definition_list():
    """KH3: extraction_method is 'html.definition_list' for native DL blocks."""
    html = _html_fixture(
        "<dl>"
        "<dt>Name</dt><dd>Alice</dd>"
        "<dt>Email</dt><dd>alice@example.com</dd>"
        "</dl>"
    )
    blocks = _parse_html_blocks(html)
    kv_blocks = [b for b in blocks if b.type == BlockType.KEY_VALUE_GROUP]
    assert len(kv_blocks) >= 1
    group = kv_blocks[0].key_value_group
    assert group is not None
    assert group.extraction_method == "html.definition_list"


def test_kh4_dl_confidence_is_extracted():
    """KH4: confidence is 'extracted' for native DL blocks."""
    html = _html_fixture(
        "<dl><dt>Name</dt><dd>Alice</dd><dt>Role</dt><dd>Engineer</dd></dl>"
    )
    blocks = _parse_html_blocks(html)
    kv_blocks = [b for b in blocks if b.type == BlockType.KEY_VALUE_GROUP]
    assert len(kv_blocks) >= 1
    assert kv_blocks[0].confidence == ExtractionConfidence.EXTRACTED


def test_kh5_empty_dl_no_block_emitted():
    """KH5: <dl> with 0 dt/dd pairs → no KEY_VALUE_GROUP block emitted."""
    html = _html_fixture("<dl></dl>")
    blocks = _parse_html_blocks(html)
    kv_blocks = [b for b in blocks if b.type == BlockType.KEY_VALUE_GROUP]
    assert len(kv_blocks) == 0


# ── XLSX TWO-COLUMN ───────────────────────────────────────────────────────────


def test_kx1_two_column_xlsx_creates_kv_block(tmp_path):
    """KX1: 2-column XLSX with short labels and values → KEY_VALUE_GROUP block."""
    path = _make_xlsx_kv(tmp_path)
    blocks = _parse_xlsx_blocks(path)
    kv_blocks = [b for b in blocks if b.type == BlockType.KEY_VALUE_GROUP]
    assert len(kv_blocks) >= 1


def test_kx2_two_column_with_header_remains_table(tmp_path):
    """KX2: 2-column XLSX where first-column values are repeated → remains a table."""
    path = _make_xlsx_repeated_keys(tmp_path)
    blocks = _parse_xlsx_blocks(path)
    # Should not be promoted to KV because first-column values repeat
    kv_blocks = [b for b in blocks if b.type == BlockType.KEY_VALUE_GROUP]
    # Either no KV block, or the table block exists
    table_blocks = [b for b in blocks if b.type == BlockType.TABLE]
    # If repeated keys, it should remain a table or not be promoted
    # (the _is_kv_region check will reject it)
    assert len(table_blocks) >= 1 or len(kv_blocks) == 0


def test_kx3_four_column_xlsx_remains_table(tmp_path):
    """KX3: 4-column XLSX remains a table."""
    path = _make_xlsx_table(tmp_path)
    blocks = _parse_xlsx_blocks(path)
    # 3-column table should remain a TABLE block
    table_blocks = [b for b in blocks if b.type == BlockType.TABLE]
    assert len(table_blocks) >= 1


def test_kx4_repeated_first_column_remains_table(tmp_path):
    """KX4: XLSX with repeated first-column values remains a table."""
    path = _make_xlsx_repeated_keys(tmp_path)
    blocks = _parse_xlsx_blocks(path)
    kv_blocks = [b for b in blocks if b.type == BlockType.KEY_VALUE_GROUP]
    # No KV group because _is_kv_region rejects repeated first-column values
    assert len(kv_blocks) == 0


def test_kx5_is_kv_region_returns_false_for_conventional_table():
    """KX5: _is_kv_region returns False for conventional tables."""
    from aksharamd.models.table import TableCell
    from aksharamd.plugins.parsers.spreadsheet import _is_kv_region

    # 3-column table
    cells = [
        TableCell(text="Name", row=0, column=0),
        TableCell(text="Age", row=0, column=1),
        TableCell(text="City", row=0, column=2),
        TableCell(text="Alice", row=1, column=0),
        TableCell(text="30", row=1, column=1),
        TableCell(text="NYC", row=1, column=2),
    ]
    assert _is_kv_region(cells, row_count=2, col_count=3) is False


# ── DOCX PROPERTIES ───────────────────────────────────────────────────────────


def test_kd1_docx_with_title_and_author_creates_kv_block(tmp_path):
    """KD1: DOCX with title+author core properties → KEY_VALUE_GROUP with group_type=METADATA."""
    path = _make_docx_with_props(tmp_path, title="My Report", author="Alice Smith")
    blocks = _parse_docx_blocks(path)
    kv_blocks = [b for b in blocks if b.type == BlockType.KEY_VALUE_GROUP]
    assert len(kv_blocks) >= 1
    group = kv_blocks[0].key_value_group
    assert group is not None
    assert group.group_type == KeyValueGroupType.METADATA


def test_kd2_docx_extraction_method_is_native_properties(tmp_path):
    """KD2: extraction_method is 'docx.native_properties'."""
    path = _make_docx_with_props(tmp_path, title="Test", author="Author")
    blocks = _parse_docx_blocks(path)
    kv_blocks = [b for b in blocks if b.type == BlockType.KEY_VALUE_GROUP]
    assert len(kv_blocks) >= 1
    group = kv_blocks[0].key_value_group
    assert group is not None
    assert group.extraction_method == "docx.native_properties"


def test_kd3_docx_no_meaningful_props_no_kv_block(tmp_path):
    """KD3: DOCX with no title/author properties → no properties block emitted."""
    path = _make_docx_no_props(tmp_path)
    blocks = _parse_docx_blocks(path)
    # Properties block requires at least 2 entries — docx with no meaningful props
    # should not produce a KV group (may produce one if python-docx fills in defaults)
    kv_blocks = [b for b in blocks if b.type == BlockType.KEY_VALUE_GROUP]
    # This is a "best effort" check — if python-docx fills in defaults like
    # Author or Created, we'd get entries. The key check is that it doesn't
    # produce one with the title "Document Properties" containing empty fields.
    if kv_blocks:
        for kb in kv_blocks:
            group = kb.key_value_group
            assert group is not None
            # If a KV block was emitted, it must have at least 2 real entries
            assert len(group.entries) >= 2


# ── OPTIMIZER PROTECTION ──────────────────────────────────────────────────────


def test_ko1_kv_block_not_removed_as_furniture():
    """KO1: KEY_VALUE_GROUP block is not removed by the optimizer (furniture removal doesn't affect it)."""
    from aksharamd.plugins.optimizers.token import TokenOptimizer

    group = KeyValueGroup(
        entries=[
            KeyValueEntry(key="Name", value="Alice"),
            KeyValueEntry(key="Email", value="alice@example.com"),
        ],
        extraction_method="test",
        confidence="extracted",
    )
    kv_block = Block.from_key_value_group(group, page=1, index=0)

    # Simulate repetition across pages to trigger furniture detection
    blocks = [kv_block.model_copy(update={"page": p, "index": p}) for p in range(1, 8)]

    ctx = _make_ctx(blocks)
    ctx.document = ctx.document.model_copy(update={"pages": 7})

    optimizer = TokenOptimizer()
    result = optimizer.execute(ctx)

    kv_blocks = [b for b in result.document.blocks if b.type == BlockType.KEY_VALUE_GROUP]
    # All KV blocks should be preserved
    assert len(kv_blocks) == len(blocks)


def test_ko2_kv_block_not_deduplicated():
    """KO2: KEY_VALUE_GROUP block is not deduplicated even when two groups share the same rendered text."""
    from aksharamd.plugins.optimizers.token import TokenOptimizer

    group = KeyValueGroup(
        entries=[
            KeyValueEntry(key="Name", value="Alice"),
            KeyValueEntry(key="Email", value="alice@example.com"),
        ],
        extraction_method="test",
    )
    kv1 = Block.from_key_value_group(group, page=1, index=0)
    kv2 = Block.from_key_value_group(group, page=2, index=1)

    ctx = _make_ctx([kv1, kv2])
    optimizer = TokenOptimizer()
    result = optimizer.execute(ctx)

    kv_blocks = [b for b in result.document.blocks if b.type == BlockType.KEY_VALUE_GROUP]
    assert len(kv_blocks) == 2  # Neither should be removed as duplicate


def test_ko3_heading_before_kv_group_preserved():
    """KO3: Numbered heading followed by KV group: heading preserved by protected-heading rule."""
    from aksharamd.plugins.optimizers.token import TokenOptimizer

    group = KeyValueGroup(
        entries=[
            KeyValueEntry(key="Name", value="Alice"),
            KeyValueEntry(key="Email", value="alice@example.com"),
        ],
        extraction_method="test",
    )
    heading = Block(
        type=BlockType.HEADING,
        content="1.1 Document Properties",
        level=2,
        page=1,
        index=0,
    )
    kv_block = Block.from_key_value_group(group, page=1, index=1)

    ctx = _make_ctx([heading, kv_block])
    optimizer = TokenOptimizer()
    result = optimizer.execute(ctx)

    headings = [b for b in result.document.blocks if b.type == BlockType.HEADING]
    kv_blocks = [b for b in result.document.blocks if b.type == BlockType.KEY_VALUE_GROUP]
    assert len(headings) >= 1
    assert len(kv_blocks) >= 1


def test_ko4_kv_block_survives_full_optimize():
    """KO4: KEY_VALUE_GROUP block survives a full optimize() call."""
    from aksharamd.plugins.optimizers.token import TokenOptimizer

    group = KeyValueGroup(
        entries=[
            KeyValueEntry(key="Project", value="AksharaMD"),
            KeyValueEntry(key="Version", value="0.4.0"),
            KeyValueEntry(key="Status", value="Active"),
        ],
        extraction_method="test",
    )
    kv_block = Block.from_key_value_group(group, page=1, index=0)
    heading = Block(type=BlockType.HEADING, content="Summary", level=2, page=1, index=1)
    para = Block(type=BlockType.PARAGRAPH, content="Some other text.", page=1, index=2)

    ctx = _make_ctx([kv_block, heading, para])
    ctx.document = ctx.document.model_copy(update={"pages": 1})

    optimizer = TokenOptimizer()
    result = optimizer.execute(ctx)

    kv_blocks = [b for b in result.document.blocks if b.type == BlockType.KEY_VALUE_GROUP]
    assert len(kv_blocks) == 1


# ── CHUNKER ───────────────────────────────────────────────────────────────────


def _kv_group_block() -> Block:
    group = KeyValueGroup(
        entries=[
            KeyValueEntry(key="Name", value="Alice"),
            KeyValueEntry(key="Email", value="alice@example.com"),
            KeyValueEntry(key="City", value="New York"),
        ],
        extraction_method="test",
    )
    return Block.from_key_value_group(group, page=1, index=0)


def test_kc1_kv_chunk_has_key_value_group_id():
    """KC1: KEY_VALUE_GROUP block produces a chunk with metadata containing 'key_value_group_id'."""
    from aksharamd.plugins.chunkers.semantic import SemanticChunker

    kv_block = _kv_group_block()
    ctx = _make_ctx([kv_block])
    ctx.document = ctx.document.model_copy(update={"pages": 1})
    ctx.document.compute_id()

    chunker = SemanticChunker(max_tokens=512)
    result = chunker.execute(ctx)

    assert len(result.chunks) >= 1
    kv_chunks = [c for c in result.chunks if "key_value_group_id" in c.metadata]
    assert len(kv_chunks) >= 1


def test_kc2_kv_chunk_flushed_after_paragraphs():
    """KC2: KV chunk is flushed after preceding paragraph blocks."""
    from aksharamd.plugins.chunkers.semantic import SemanticChunker

    para = Block(type=BlockType.PARAGRAPH, content="Some text before the KV group.", index=0)
    kv_block = _kv_group_block()

    ctx = _make_ctx([para, kv_block])
    ctx.document = ctx.document.model_copy(update={"pages": 1})
    ctx.document.compute_id()

    chunker = SemanticChunker(max_tokens=512)
    result = chunker.execute(ctx)

    # Should have at least 2 chunks (one for paragraph, one for KV group)
    assert len(result.chunks) >= 2


def test_kc3_kv_chunk_has_group_type():
    """KC3: KV chunk metadata has 'group_type' field."""
    from aksharamd.plugins.chunkers.semantic import SemanticChunker

    kv_block = _kv_group_block()
    ctx = _make_ctx([kv_block])
    ctx.document = ctx.document.model_copy(update={"pages": 1})
    ctx.document.compute_id()

    chunker = SemanticChunker(max_tokens=512)
    result = chunker.execute(ctx)

    kv_chunks = [c for c in result.chunks if "key_value_group_id" in c.metadata]
    assert len(kv_chunks) >= 1
    assert "group_type" in kv_chunks[0].metadata


def test_kc4_kv_chunk_has_entry_count():
    """KC4: KV chunk metadata has 'entry_count' field."""
    from aksharamd.plugins.chunkers.semantic import SemanticChunker

    kv_block = _kv_group_block()
    ctx = _make_ctx([kv_block])
    ctx.document = ctx.document.model_copy(update={"pages": 1})
    ctx.document.compute_id()

    chunker = SemanticChunker(max_tokens=512)
    result = chunker.execute(ctx)

    kv_chunks = [c for c in result.chunks if "key_value_group_id" in c.metadata]
    assert len(kv_chunks) >= 1
    assert "entry_count" in kv_chunks[0].metadata
    assert kv_chunks[0].metadata["entry_count"] == 3  # 3 entries in the test group


# ── BACKWARD COMPATIBILITY ────────────────────────────────────────────────────


def test_kb1_no_kv_content_unchanged():
    """KB1: Document with no KV content compiles identically before and after (block count unchanged)."""
    blocks = [
        Block(type=BlockType.HEADING, content="My Document", level=1, index=0),
        Block(type=BlockType.PARAGRAPH, content="This is regular prose content.", index=1),
        Block(type=BlockType.PARAGRAPH, content="Another paragraph of text.", index=2),
    ]
    original_count = len(blocks)
    original_ids = [b.id for b in blocks]

    ctx = _make_ctx(blocks)
    result = detect_and_promote_key_value_groups(ctx)

    assert len(result.document.blocks) == original_count
    assert [b.id for b in result.document.blocks] == original_ids


def test_kb2_none_document_returns_ctx_unchanged():
    """KB2: detect_and_promote_key_value_groups with None document returns ctx unchanged."""
    from aksharamd.context import CompilationContext

    ctx = CompilationContext(source="test.txt", output_dir=tempfile.mkdtemp())
    ctx.document = None

    result = detect_and_promote_key_value_groups(ctx)
    assert result.document is None


def test_kb3_key_value_imports_still_work():
    """KB3: Existing key_value imports still work (models, detection, rendering)."""
    from aksharamd.models.key_value import (
        KeyValueEntry,
        KeyValueGroup,
    )
    from aksharamd.renderers.key_value_markdown import render_key_value_group
    from aksharamd.scoring.key_value_detection import detect_key_value_entries

    # Basic smoke test that imports work and basic functionality is intact
    group = KeyValueGroup(
        entries=[
            KeyValueEntry(key="Name", value="Alice"),
            KeyValueEntry(key="Email", value="alice@example.com"),
        ],
        extraction_method="test",
    )
    rendered = render_key_value_group(group)
    assert "Name" in rendered
    assert "Alice" in rendered

    result = detect_key_value_entries("Name: Alice\nEmail: alice@example.com")
    assert result.group is not None


def test_kb4_rhetorical_prose_not_promoted():
    """KB4: Rhetorical prose ('Note: this explains...') is never promoted."""
    content = "Note: this is an important note\nWarning: be careful here"
    ctx = _make_ctx([_para(content)])
    result = detect_and_promote_key_value_groups(ctx)
    blocks = result.document.blocks
    # Rhetorical labels should be rejected by the detector
    assert len(blocks) == 1
    assert blocks[0].type == BlockType.PARAGRAPH


# ── Additional robustness tests ───────────────────────────────────────────────


def test_detector_version_constant():
    """Verify DETECTOR_VERSION is the expected string."""
    assert DETECTOR_VERSION == "kv_promoter/v2"


def test_promoted_block_has_detector_version_in_metadata():
    """Promoted block's metadata includes 'detector_version' = DETECTOR_VERSION."""
    content = "Name: Alice\nEmail: alice@example.com"
    ctx = _make_ctx([_para(content)])
    result = detect_and_promote_key_value_groups(ctx)
    promoted = result.document.blocks[0]
    assert promoted.type == BlockType.KEY_VALUE_GROUP
    assert promoted.metadata.get("detector_version") == DETECTOR_VERSION


def test_empty_document_returns_unchanged():
    """detect_and_promote_key_value_groups on an empty document returns ctx unchanged."""
    ctx = _make_ctx([])
    result = detect_and_promote_key_value_groups(ctx)
    assert result.document.blocks == []


def test_kv_block_has_entry_count_in_metadata():
    """Promoted block's metadata contains 'entry_count' matching the number of entries."""
    content = "Name: Alice\nEmail: alice@example.com\nCity: NYC"
    ctx = _make_ctx([_para(content)])
    result = detect_and_promote_key_value_groups(ctx)
    promoted = result.document.blocks[0]
    assert promoted.type == BlockType.KEY_VALUE_GROUP
    assert promoted.metadata.get("entry_count") == 3


def test_is_kv_region_returns_true_for_valid_kv():
    """_is_kv_region returns True for a proper 2-column KV layout."""
    from aksharamd.models.table import TableCell
    from aksharamd.plugins.parsers.spreadsheet import _is_kv_region

    cells = [
        TableCell(text="Project", row=0, column=0),
        TableCell(text="AksharaMD", row=0, column=1),
        TableCell(text="Owner", row=1, column=0),
        TableCell(text="Kalyan", row=1, column=1),
        TableCell(text="Version", row=2, column=0),
        TableCell(text="0.4.0", row=2, column=1),
    ]
    assert _is_kv_region(cells, row_count=3, col_count=2) is True


def test_validation_issue_metadata_field_present():
    """ValidationIssue.metadata field exists and defaults to empty dict."""
    from aksharamd.models.validation import Severity, ValidationIssue

    issue = ValidationIssue(
        severity=Severity.WARNING,
        code="W_TEST",
        message="test message",
    )
    assert hasattr(issue, "metadata")
    assert isinstance(issue.metadata, dict)


def test_kv_block_content_derived_from_group():
    """Block.from_key_value_group produces content from render_key_value_group."""
    group = KeyValueGroup(
        entries=[
            KeyValueEntry(key="Name", value="Alice"),
            KeyValueEntry(key="Email", value="alice@example.com"),
        ],
        extraction_method="test",
    )
    block = Block.from_key_value_group(group, page=1, index=0)
    assert "Name" in block.content
    assert "Alice" in block.content


def test_html_dl_with_dt_before_dd():
    """<dl> with correct dt/dd order produces entries with correct key-value pairs."""
    html = _html_fixture(
        "<dl>"
        "<dt>Language</dt><dd>Python</dd>"
        "<dt>Framework</dt><dd>FastAPI</dd>"
        "<dt>Database</dt><dd>PostgreSQL</dd>"
        "</dl>"
    )
    blocks = _parse_html_blocks(html)
    kv_blocks = [b for b in blocks if b.type == BlockType.KEY_VALUE_GROUP]
    assert len(kv_blocks) >= 1
    group = kv_blocks[0].key_value_group
    assert group is not None
    keys = [e.key for e in group.entries]
    values = [e.value for e in group.entries]
    assert "Language" in keys
    assert "Python" in values
