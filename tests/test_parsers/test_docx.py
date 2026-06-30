from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from docx.oxml.ns import qn as qname
from lxml import etree

from aksharamd.context import CompilationContext
from aksharamd.models.block import BlockType
from aksharamd.plugins.parsers.docx import DocxParser


def _make_docx(tmp_path: Path, builder) -> Path:
    doc = DocxDocument()
    builder(doc)
    p = tmp_path / "test.docx"
    doc.save(str(p))
    return p


def _parse(path: Path, tmp_path: Path) -> CompilationContext:
    ctx = CompilationContext(source=str(path), output_dir=str(tmp_path / "out"))
    return DocxParser().execute(ctx)


def test_heading_styles(tmp_path):
    path = _make_docx(tmp_path, lambda doc: [
        doc.add_heading("Main Title", level=1),
        doc.add_heading("Chapter", level=2),
        doc.add_heading("Section", level=3),
    ])
    ctx = _parse(path, tmp_path)
    headings = [b for b in ctx.document.blocks if b.type == BlockType.HEADING]
    assert headings[0].level == 1 and headings[0].content == "Main Title"
    assert headings[1].level == 2
    assert headings[2].level == 3


def test_paragraph(tmp_path):
    path = _make_docx(tmp_path, lambda doc: doc.add_paragraph("Hello from DOCX."))
    ctx = _parse(path, tmp_path)
    paras = [b for b in ctx.document.blocks if b.type == BlockType.PARAGRAPH]
    assert any("Hello from DOCX" in b.content for b in paras)


def test_bullet_list_style(tmp_path):
    def build(doc):
        doc.add_paragraph("Alpha", style="List Bullet")
        doc.add_paragraph("Beta", style="List Bullet")
        doc.add_paragraph("Gamma", style="List Bullet")
    path = _make_docx(tmp_path, build)
    ctx = _parse(path, tmp_path)
    lists = [b for b in ctx.document.blocks if b.type == BlockType.LIST]
    assert len(lists) == 1
    assert "- Alpha" in lists[0].content
    assert "- Gamma" in lists[0].content


def test_nested_list_style(tmp_path):
    def build(doc):
        doc.add_paragraph("Top A", style="List Bullet")
        doc.add_paragraph("Nested B", style="List Bullet 2")
        doc.add_paragraph("Deep C", style="List Bullet 3")
        doc.add_paragraph("Top D", style="List Bullet")
    path = _make_docx(tmp_path, build)
    ctx = _parse(path, tmp_path)
    lists = [b for b in ctx.document.blocks if b.type == BlockType.LIST]
    assert len(lists) == 1
    content = lists[0].content
    assert "- Top A" in content
    assert "  - Nested B" in content
    assert "    - Deep C" in content


def test_numbered_list_style(tmp_path):
    def build(doc):
        doc.add_paragraph("Step one", style="List Number")
        doc.add_paragraph("Step two", style="List Number")
        doc.add_paragraph("Step three", style="List Number")
    path = _make_docx(tmp_path, build)
    ctx = _parse(path, tmp_path)
    lists = [b for b in ctx.document.blocks if b.type == BlockType.LIST]
    assert len(lists) == 1
    assert "1. Step one" in lists[0].content
    assert "3. Step three" in lists[0].content


def test_list_between_paragraphs(tmp_path):
    def build(doc):
        doc.add_paragraph("Before.")
        doc.add_paragraph("Item A", style="List Bullet")
        doc.add_paragraph("Item B", style="List Bullet")
        doc.add_paragraph("After.")
    path = _make_docx(tmp_path, build)
    ctx = _parse(path, tmp_path)
    types = [b.type for b in ctx.document.blocks]
    assert BlockType.PARAGRAPH in types
    assert BlockType.LIST in types
    # List must be between the two paragraphs
    para_indices = [i for i, b in enumerate(ctx.document.blocks) if b.type == BlockType.PARAGRAPH]
    list_index = next(i for i, b in enumerate(ctx.document.blocks) if b.type == BlockType.LIST)
    assert para_indices[0] < list_index < para_indices[-1]


def test_two_separate_lists(tmp_path):
    def build(doc):
        doc.add_paragraph("Bullet A", style="List Bullet")
        doc.add_paragraph("Bullet B", style="List Bullet")
        doc.add_paragraph("Between.")
        doc.add_paragraph("Number 1", style="List Number")
        doc.add_paragraph("Number 2", style="List Number")
    path = _make_docx(tmp_path, build)
    ctx = _parse(path, tmp_path)
    lists = [b for b in ctx.document.blocks if b.type == BlockType.LIST]
    assert len(lists) == 2


def test_table(tmp_path):
    def build(doc):
        table = doc.add_table(rows=2, cols=3)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Age"
        table.cell(0, 2).text = "City"
        table.cell(1, 0).text = "Alice"
        table.cell(1, 1).text = "30"
        table.cell(1, 2).text = "London"
    path = _make_docx(tmp_path, build)
    ctx = _parse(path, tmp_path)
    tables = [b for b in ctx.document.blocks if b.type == BlockType.TABLE]
    assert len(tables) == 1
    assert "Name" in tables[0].content and "Alice" in tables[0].content


def test_numpr_list_detection(tmp_path):
    """Verify w:numPr XML-based list detection works (real Word document format)."""
    doc = DocxDocument()

    def add_numpr_para(text, num_id, ilvl):
        p = doc.add_paragraph()
        pPr = p._element.get_or_add_pPr()
        numPr = etree.SubElement(pPr, qname("w:numPr"))
        ilvl_el = etree.SubElement(numPr, qname("w:ilvl"))
        ilvl_el.set(qname("w:val"), str(ilvl))
        numId_el = etree.SubElement(numPr, qname("w:numId"))
        numId_el.set(qname("w:val"), str(num_id))
        p.add_run(text)

    add_numpr_para("Item 1", 1, 0)
    add_numpr_para("Item 2", 1, 0)
    add_numpr_para("Nested", 1, 1)
    add_numpr_para("Item 3", 1, 0)

    path = tmp_path / "numpr.docx"
    doc.save(str(path))
    ctx = _parse(path, tmp_path)

    lists = [b for b in ctx.document.blocks if b.type == BlockType.LIST]
    assert len(lists) == 1
    content = lists[0].content
    assert "Item 1" in content
    assert "  - Nested" in content


def test_corrupt_file_returns_error(tmp_path):
    p = tmp_path / "corrupt.docx"
    p.write_bytes(b"not a real docx file")
    ctx = _parse(p, tmp_path)
    assert ctx.document is None
    assert any(e.code == "DOCX_PARSE_ERROR" for e in ctx.validation.errors)
