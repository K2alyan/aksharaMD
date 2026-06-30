from __future__ import annotations

from pathlib import Path

from lxml import etree

from aksharamd.context import CompilationContext
from aksharamd.models.block import BlockType
from aksharamd.plugins.parsers.docx import _omml_to_latex

_M = "http://schemas.openxmlformats.org/officeDocument/2006/math"
_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ── Builder helpers ──────────────────────────────────────────────────────────

def m(tag, *children, text=None):
    """Build an OMML element in the math namespace."""
    el = etree.Element(f"{{{_M}}}{tag}")
    if text is not None:
        el.text = text
    for c in children:
        el.append(c)
    return el


def mr(text: str):
    """Shorthand for m:r containing m:t with text."""
    return m("r", m("t", text=text))


# ── _omml_to_latex unit tests ────────────────────────────────────────────────

def test_plain_text():
    el = m("oMath", mr("x"))
    assert _omml_to_latex(el) == "x"


def test_superscript():
    el = m("sSup", m("e", mr("x")), m("sup", mr("2")))
    assert _omml_to_latex(el) == "{x}^{2}"


def test_subscript():
    el = m("sSub", m("e", mr("a")), m("sub", mr("i")))
    assert _omml_to_latex(el) == "{a}_{i}"


def test_sub_superscript():
    el = m("sSubSup", m("e", mr("x")), m("sub", mr("0")), m("sup", mr("n")))
    assert _omml_to_latex(el) == "{x}_{0}^{n}"


def test_fraction():
    el = m("f", m("num", mr("1")), m("den", mr("2")))
    assert _omml_to_latex(el) == "\\frac{1}{2}"


def test_sqrt_no_degree():
    el = m("rad", m("deg"), m("e", mr("x")))
    result = _omml_to_latex(el)
    assert result == "\\sqrt{x}"


def test_sqrt_with_degree():
    el = m("rad", m("deg", mr("3")), m("e", mr("x")))
    result = _omml_to_latex(el)
    assert result == "\\sqrt[3]{x}"


def test_delimiter():
    el = m("d", m("e", mr("x+y")))
    result = _omml_to_latex(el)
    assert "\\left(" in result and "\\right)" in result
    assert "x+y" in result


def test_nested_fraction():
    # (1/2)^2
    frac = m("f", m("num", mr("1")), m("den", mr("2")))
    el = m("sSup", m("e", frac), m("sup", mr("2")))
    result = _omml_to_latex(el)
    assert "\\frac" in result and "^{2}" in result


def test_fallback_collects_leaf_text():
    # Unknown tag — should still collect m:t text
    inner = m("unknownTag", m("r", m("t", text="hello")))
    result = _omml_to_latex(inner)
    assert "hello" in result


# ── Integration test: DOCX with math paragraphs ──────────────────────────────

def _make_math_docx(tmp_path: Path) -> Path:
    """Create a DOCX file with a block equation using raw lxml injection."""
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_paragraph("Before equation.")

    # Inject a w:oMathPara containing m:oMath with x^2

    body = doc.element.body
    oMathPara = etree.SubElement(body, f"{{{_M}}}oMathPara")
    oMath = etree.SubElement(oMathPara, f"{{{_M}}}oMath")
    sSup = etree.SubElement(oMath, f"{{{_M}}}sSup")
    e_el = etree.SubElement(sSup, f"{{{_M}}}e")
    r_el = etree.SubElement(e_el, f"{{{_M}}}r")
    t_el = etree.SubElement(r_el, f"{{{_M}}}t")
    t_el.text = "x"
    sup_el = etree.SubElement(sSup, f"{{{_M}}}sup")
    r2 = etree.SubElement(sup_el, f"{{{_M}}}r")
    t2 = etree.SubElement(r2, f"{{{_M}}}t")
    t2.text = "2"

    doc.add_paragraph("After equation.")
    path = tmp_path / "math.docx"
    doc.save(str(path))
    return path


def test_block_equation_renders_as_latex(tmp_path):
    path = _make_math_docx(tmp_path)
    from aksharamd.plugins.parsers.docx import DocxParser

    ctx = CompilationContext(source=str(path), output_dir=str(tmp_path / "out"))
    ctx = DocxParser().execute(ctx)

    assert ctx.document is not None
    math_blocks = [b for b in ctx.document.blocks if "$$" in b.content]
    assert len(math_blocks) >= 1
    content = math_blocks[0].content
    assert "x" in content
    assert "^" in content or "2" in content


def test_non_math_paragraphs_unaffected(tmp_path):
    path = _make_math_docx(tmp_path)
    from aksharamd.plugins.parsers.docx import DocxParser

    ctx = CompilationContext(source=str(path), output_dir=str(tmp_path / "out"))
    ctx = DocxParser().execute(ctx)

    paras = [b for b in ctx.document.blocks if b.type == BlockType.PARAGRAPH]
    plain = [p for p in paras if "$$" not in p.content]
    assert any("Before equation" in p.content for p in plain)
    assert any("After equation" in p.content for p in plain)
