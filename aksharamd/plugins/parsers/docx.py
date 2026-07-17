from __future__ import annotations

import os
import re
from itertools import groupby as _groupby
from pathlib import Path

from ...context import CompilationContext
from ...models.asset import Asset
from ...models.block import Block, BlockType, ExtractionConfidence
from ...models.document import Document
from ...models.table import ExtractionMethod, TableCell, TableData
from ..base import ParserPlugin
from ..registry import register_parser

_DRAW_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
_REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

_LEVEL_SUFFIX_RE = re.compile(r'\s+(\d+)$')
_W_BR = f"{{{_W_NS}}}br"
_W_TYPE = f"{{{_W_NS}}}type"
# Cap paragraph count to prevent memory exhaustion on adversarially large documents.
_MAX_DOCX_PARAGRAPHS = int(os.environ.get("AKSHARAMD_MAX_DOCX_PARAGRAPHS", "10000"))


def _has_page_break(para_el) -> bool:
    """Return True if this paragraph element contains an explicit page break."""
    for br in para_el.iter(_W_BR):
        if br.get(_W_TYPE) == "page":
            return True
    return False


def _extract_drawing_bytes(para_el, doc_part) -> list[bytes]:
    """Extract image bytes from all w:drawing elements within a paragraph element."""
    images = []
    for drawing in para_el.iter(f"{{{_W_NS}}}drawing"):
        for blip in drawing.iter(f"{{{_DRAW_NS}}}blip"):
            r_id = blip.get(f"{{{_REL_NS}}}embed")
            if r_id:
                try:
                    images.append(doc_part.related_parts[r_id].blob)
                except Exception:
                    pass
    return images

_OMML_URI = "http://schemas.openxmlformats.org/officeDocument/2006/math"
_M = f"{{{_OMML_URI}}}"


def _omml_to_latex(el) -> str:
    """Recursively convert an OMML element subtree to a LaTeX string."""
    tag = el.tag[len(_M):] if el.tag.startswith(_M) else el.tag.split("}")[-1]

    def sub(child_tag: str) -> str:
        child = el.find(f"{_M}{child_tag}")
        return _omml_to_latex(child) if child is not None else ""

    def children() -> str:
        return "".join(_omml_to_latex(c) for c in el)

    if tag == "t":
        return el.text or ""
    if tag in ("r", "e", "oMath"):
        return children()
    if tag == "sSup":
        return f"{{{sub('e')}}}^{{{sub('sup')}}}"
    if tag == "sSub":
        return f"{{{sub('e')}}}_{{{sub('sub')}}}"
    if tag == "sSubSup":
        return f"{{{sub('e')}}}_{{{sub('sub')}}}^{{{sub('sup')}}}"
    if tag == "f":
        return f"\\frac{{{sub('num')}}}{{{sub('den')}}}"
    if tag == "rad":
        deg = sub("deg")
        return f"\\sqrt[{deg}]{{{sub('e')}}}" if deg.strip() else f"\\sqrt{{{sub('e')}}}"
    if tag == "d":
        return f"\\left({sub('e')}\\right)"
    if tag == "nary":
        chr_el = el.find(f"{_M}naryPr/{_M}chr")
        op = chr_el.text if chr_el is not None and chr_el.text else "\\int"
        return f"{op}_{{{sub('sub')}}}^{{{sub('sup')}}}{{{sub('e')}}}"
    if tag == "m":  # matrix
        rows = [
            " & ".join(_omml_to_latex(c) for c in r.findall(f"{_M}e"))
            for r in el.findall(f"{_M}mr")
        ]
        return "\\begin{pmatrix}" + " \\\\ ".join(rows) + "\\end{pmatrix}"
    # Fallback: collect all leaf m:t text
    parts = [n.text for n in el.iter(f"{_M}t") if n.text]
    return " ".join(parts)


def _get_list_props(para, qn) -> tuple[str | None, int, bool | None]:
    """
    Return (group_key, ilvl, is_ordered) if this paragraph is a list item, else (None, 0, False).

    group_key  — identifies which list this belongs to (for grouping consecutive items)
    ilvl       — indent level 0=top
    is_ordered — True=numbered, False=bullet, None=must look up from numbering XML
    """
    # Primary: w:numPr in paragraph XML (authoritative for real Word documents)
    try:
        pPr = para._element.find(qn("w:pPr"))
        if pPr is not None:
            numPr = pPr.find(qn("w:numPr"))
            if numPr is not None:
                numId_el = numPr.find(qn("w:numId"))
                ilvl_el = numPr.find(qn("w:ilvl"))
                if numId_el is not None:
                    numId = numId_el.get(qn("w:val"), "0")
                    if numId != "0":
                        ilvl = int(ilvl_el.get(qn("w:val"), "0")) if ilvl_el is not None else 0
                        return f"num:{numId}", ilvl, None  # is_ordered resolved from numbering XML
    except Exception:
        pass

    # Fallback: style name (python-docx-created files and simple Word docs)
    style_name = (para.style.name or "") if para.style else ""
    sl = style_name.lower()
    if "list bullet" in sl:
        m = _LEVEL_SUFFIX_RE.search(style_name)
        ilvl = int(m.group(1)) - 1 if m else 0
        return "style:bullet", ilvl, False
    if "list number" in sl:
        m = _LEVEL_SUFFIX_RE.search(style_name)
        ilvl = int(m.group(1)) - 1 if m else 0
        return "style:number", ilvl, True

    return None, 0, False


def _build_ordered_numids(doc) -> set[str]:
    """Return the set of numIds whose level-0 format is ordered (not bullet/none)."""
    ordered: set[str] = set()
    try:
        from docx.oxml.ns import qn
        root = doc.part.numbering_part._element
        # abstractNumId -> level-0 numFmt value
        abstract_fmt0: dict[str, str] = {}
        for a in root.findall(qn("w:abstractNum")):
            aid = a.get(qn("w:abstractNumId"))
            for lvl in a.findall(qn("w:lvl")):
                if lvl.get(qn("w:ilvl")) == "0":
                    nf = lvl.find(qn("w:numFmt"))
                    fmt = nf.get(qn("w:val"), "bullet") if nf is not None else "bullet"
                    abstract_fmt0[aid] = fmt
                    break
        for n in root.findall(qn("w:num")):
            nid = n.get(qn("w:numId"))
            ref = n.find(qn("w:abstractNumId"))
            if ref is None:
                continue
            fmt = abstract_fmt0.get(ref.get(qn("w:val")), "bullet")
            if fmt not in ("bullet", "none"):
                ordered.add(nid)
    except Exception:
        pass
    return ordered


_HEADING_STYLES = {
    "heading 1": 1, "heading 2": 2, "heading 3": 3,
    "heading 4": 4, "heading 5": 5, "heading 6": 6,
    "title": 1, "subtitle": 2,
}


_DOCX_CELL_TEXT_MAX = 200


def _extract_docx_table_data(table, page: int | None, qn) -> TableData:
    """Parse a python-docx Table into TableData with rowspan/colspan support."""
    rows = table.rows
    if not rows:
        return TableData(
            row_count=0, column_count=0, cells=[],
            extraction_method=ExtractionMethod.DOCX_NATIVE,
            span_detection="native",
            header_detection="assumed_first_row",
        )

    try:
        col_count = len(table.columns)
    except Exception:
        col_count = max(len(r.cells) for r in rows) if rows else 0
    row_count = len(rows)

    from collections import namedtuple
    PCel = namedtuple('PCel', ['cell', 'c', 'col_span', 'vmerge'])
    grid: list[list[PCel]] = []

    for r_idx, row in enumerate(rows):
        row_entries = []
        c_offset = 0
        # Group by identity to detect horizontally merged cells
        for _, cell_group in _groupby(row.cells, key=id):
            cell_list = list(cell_group)
            cell = cell_list[0]
            col_span = len(cell_list)

            # Detect vMerge from XML
            tc = cell._tc
            tc_pr = tc.find(qn('w:tcPr'))
            vmerge = None
            if tc_pr is not None:
                vm = tc_pr.find(qn('w:vMerge'))
                if vm is not None:
                    val = vm.get(qn('w:val'), '')
                    vmerge = 'restart' if val == 'restart' else 'slave'

            row_entries.append(PCel(cell, c_offset, col_span, vmerge))
            c_offset += col_span

        col_count = max(col_count, c_offset)
        grid.append(row_entries)

    # Build lookup: (r, c) -> PCel
    lookup: dict[tuple[int, int], PCel] = {}
    for r_idx, entries in enumerate(grid):
        for entry in entries:
            lookup[(r_idx, entry.c)] = entry

    table_cells: list[TableCell] = []
    for r_idx, entries in enumerate(grid):
        for entry in entries:
            if entry.vmerge == 'slave':
                continue

            row_span = 1
            if entry.vmerge == 'restart':
                for r in range(r_idx + 1, row_count):
                    slave = lookup.get((r, entry.c))
                    if slave is not None and slave.vmerge == 'slave':
                        row_span += 1
                    else:
                        break

            cell_text = " ".join(
                p.text.strip() for p in entry.cell.paragraphs if p.text.strip()
            )
            cell_text = cell_text[:_DOCX_CELL_TEXT_MAX]

            table_cells.append(TableCell(
                text=cell_text,
                row=r_idx,
                column=entry.c,
                row_span=row_span,
                column_span=entry.col_span,
            ))

    return TableData(
        row_count=row_count,
        column_count=col_count,
        cells=table_cells,
        header_rows=[0] if table_cells else [],
        header_detection="assumed_first_row",
        span_detection="native",
        extraction_method=ExtractionMethod.DOCX_NATIVE,
        page=page,
    )


def _extract_docx_properties(doc, idx: int) -> Block | None:
    """Extract core properties (title, author, subject, etc.) as a KeyValueGroup."""
    from ...models.key_value import KeyValueEntry, KeyValueGroup, KeyValueGroupType

    try:
        props = doc.core_properties
    except Exception:
        return None

    field_map = [
        ("Title", "title"),
        ("Author", "author"),
        ("Subject", "subject"),
        ("Description", "description"),
        ("Category", "category"),
        ("Keywords", "keywords"),
        ("Created", "created"),
        ("Modified", "modified"),
    ]

    entries = []
    for label, attr in field_map:
        try:
            val = getattr(props, attr, None)
            if val is not None:
                val_str = str(val).strip()
                if val_str and val_str not in ("None", ""):
                    entries.append(KeyValueEntry(
                        key=label,
                        value=val_str[:80],
                        confidence="extracted",
                    ))
        except Exception:
            pass

    if len(entries) < 2:
        return None

    group = KeyValueGroup(
        entries=entries,
        title="Document Properties",
        group_type=KeyValueGroupType.METADATA,
        extraction_method="docx.native_properties",
        confidence="extracted",
    )

    return Block.from_key_value_group(
        group,
        page=1,
        index=idx,
        confidence=ExtractionConfidence.EXTRACTED,
        metadata={"source": "docx_core_properties"},
    )


class DocxParser(ParserPlugin):
    name = "docx_parser"
    supported_types = ["docx"]

    def execute(self, ctx: CompilationContext) -> CompilationContext:
        from docx import Document as DocxDocument
        from docx.oxml.ns import qn

        path = Path(ctx.source)
        try:
            doc = DocxDocument(str(path))
        except Exception as e:
            ctx.error("DOCX_PARSE_ERROR", str(e))
            return ctx

        blocks: list[Block] = []
        assets: list[Asset] = []
        idx = 0
        title: str | None = None
        author: str | None = None

        props = doc.core_properties
        if props.title:
            title = props.title
        if props.author:
            author = props.author

        # Extract document core properties as a KeyValueGroup
        kv_block = _extract_docx_properties(doc, idx)
        if kv_block is not None:
            blocks.append(kv_block)
            idx += 1

        ordered_numIds = _build_ordered_numids(doc)

        body = doc.element.body
        para_map = {p._element: p for p in doc.paragraphs}
        table_map = {t._element: t for t in doc.tables}

        if len(para_map) > _MAX_DOCX_PARAGRAPHS:
            ctx.error(
                "DOCX_TOO_MANY_PARAGRAPHS",
                f"Document has {len(para_map)} paragraphs; limit is {_MAX_DOCX_PARAGRAPHS}. "
                f"Set AKSHARAMD_MAX_DOCX_PARAGRAPHS to increase the limit.",
            )
            return ctx

        # List accumulator — groups consecutive list paragraphs into one LIST block
        list_items: list[tuple[int, bool, str]] = []  # (ilvl, is_ordered, text)
        current_list_key: str | None = None

        def flush_list() -> None:
            nonlocal idx, current_list_key
            if not list_items:
                return
            counters: dict[int, int] = {}
            prev_ilvl = -1
            lines: list[str] = []
            for ilvl, is_ordered, text in list_items:
                # Reset deeper-level counters when ascending
                if ilvl < prev_ilvl:
                    for k in [k for k in counters if k > ilvl]:
                        del counters[k]
                indent = "  " * ilvl
                if is_ordered:
                    counters[ilvl] = counters.get(ilvl, 0) + 1
                    prefix = f"{counters[ilvl]}."
                else:
                    prefix = "-"
                lines.append(f"{indent}{prefix} {text}")
                prev_ilvl = ilvl
            blocks.append(Block(type=BlockType.LIST, content="\n".join(lines), index=idx,
                                page=current_page))
            idx += 1
            list_items.clear()
            current_list_key = None

        current_page = 1

        for child in body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if tag == "p" and child in para_map:
                para = para_map[child]

                # Page break detection — increment before processing this paragraph's content
                if _has_page_break(child):
                    current_page += 1

                # Inline images flush the current list (images break list continuity)
                img_list = _extract_drawing_bytes(child, doc.part)
                if img_list:
                    flush_list()
                    for img_bytes in img_list:
                        asset_id = f"img_{idx}"
                        assets.append(Asset(id=asset_id, type="image", image_bytes=img_bytes))
                        blocks.append(Block(type=BlockType.IMAGE, content="", index=idx,
                                            metadata={"asset_id": asset_id}, page=current_page))
                        idx += 1

                # Build mixed text+math content by walking child nodes in order
                mixed_parts: list[str] = []
                for node in child:
                    node_tag = node.tag.split("}")[-1] if "}" in node.tag else node.tag
                    if node_tag == "r":
                        t_nodes = node.findall(f"{{{_W_NS}}}t")
                        mixed_parts.extend(t.text for t in t_nodes if t.text)
                    elif node_tag == "oMath":
                        latex = _omml_to_latex(node).strip()
                        if latex:
                            mixed_parts.append(f"${latex}$")
                mixed_text = "".join(mixed_parts).strip()
                text = para.text.strip()
                display_text = mixed_text or text

                if not display_text:
                    continue

                has_math = any(
                    node.tag == f"{_M}oMath" for node in child
                )
                if has_math and not text:
                    # Pure math paragraph — emit as block equation
                    flush_list()
                    latex = _omml_to_latex(child).strip()
                    blocks.append(Block(type=BlockType.PARAGRAPH,
                                        content=f"$${latex}$$", index=idx))
                    idx += 1
                    continue

                if not text:
                    continue

                # List item?
                key, ilvl, is_ordered = _get_list_props(para, qn)
                if key is not None:
                    if is_ordered is None:
                        is_ordered = key[len("num:"):] in ordered_numIds
                    if key != current_list_key and current_list_key is not None:
                        flush_list()
                    current_list_key = key
                    list_items.append((ilvl, is_ordered, display_text))
                    continue

                # Regular paragraph
                flush_list()
                style_name = (para.style.name or "").lower() if para.style else ""
                level = _HEADING_STYLES.get(style_name)
                if level:
                    if not title and level == 1:
                        title = display_text
                    blocks.append(Block(type=BlockType.HEADING, content=display_text,
                                        level=level, index=idx, page=current_page))
                elif "code" in style_name or "mono" in style_name:
                    blocks.append(Block(type=BlockType.CODE_BLOCK, content=display_text,
                                        index=idx, page=current_page))
                else:
                    blocks.append(Block(type=BlockType.PARAGRAPH, content=display_text,
                                        index=idx, page=current_page))
                idx += 1

            elif tag == "oMathPara":
                flush_list()
                latex = _omml_to_latex(child).strip()
                if latex:
                    blocks.append(Block(type=BlockType.PARAGRAPH,
                                        content=f"$${latex}$$", index=idx, page=current_page))
                    idx += 1

            elif tag == "tbl" and child in table_map:
                flush_list()
                table = table_map[child]
                table_data = _extract_docx_table_data(table, current_page, qn)
                if table_data.cells:
                    blocks.append(Block.from_table(table_data, page=current_page, index=idx))
                    idx += 1

        flush_list()

        image_count = len([b for b in blocks if b.type == BlockType.IMAGE])
        ctx.document = Document(
            source=str(path),
            file_type="docx",
            title=title,
            author=author,
            pages=len(doc.sections),
            blocks=blocks,
            assets=assets,
            metadata={"images": image_count, "sections": len(doc.sections)},
        ).compute_id()
        return ctx


register_parser("docx", DocxParser)
